#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HTML转Word转换器
用于将实习报告HTML文件转换为Word文档格式

依赖包安装:
pip install python-docx beautifulsoup4 lxml

使用方法:
python html_to_word_converter.py
"""

import os
import sys
from pathlib import Path
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn


class HTMLToWordConverter:
    def __init__(self, html_file_path, output_file_path=None):
        """
        初始化转换器
        
        Args:
            html_file_path (str): HTML文件路径
            output_file_path (str): 输出Word文件路径，如果为None则自动生成
        """
        self.html_file_path = Path(html_file_path)
        if output_file_path is None:
            self.output_file_path = self.html_file_path.parent / f"{self.html_file_path.stem}.docx"
        else:
            self.output_file_path = Path(output_file_path)
        
        # 创建Word文档
        self.doc = Document()
        self.setup_document_styles()
    
    def setup_document_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Inches(0.5)  # 首行缩进2字符
        
        # 创建标题样式
        for i in range(1, 5):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = '宋体'
            heading_style.font.bold = True
            heading_style.font.size = Pt(14 - i)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    def clean_text(self, text):
        """清理文本内容"""
        if not text:
            return ""
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text.strip())
        return text
    
    def add_heading(self, text, level=1):
        """添加标题"""
        text = self.clean_text(text)
        if text:
            heading = self.doc.add_heading(text, level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_paragraph(self, text, style=None):
        """添加段落"""
        text = self.clean_text(text)
        if text:
            paragraph = self.doc.add_paragraph(text, style)
            return paragraph
        return None
    
    def add_list_item(self, text, is_numbered=False):
        """添加列表项"""
        text = self.clean_text(text)
        if text:
            # 移除HTML强调标签
            text = re.sub(r'<[^>]+>', '', text)
            paragraph = self.doc.add_paragraph(text, style='List Bullet')
            return paragraph
        return None
    
    def add_code_block(self, code_text):
        """添加代码块"""
        code_text = self.clean_text(code_text)
        if code_text:
            # 解码HTML实体
            code_text = code_text.replace('&lt;', '<').replace('&gt;', '>')
            code_text = code_text.replace('&amp;', '&')
            
            paragraph = self.doc.add_paragraph(code_text)
            paragraph.style.font.name = 'Courier New'
            paragraph.style.font.size = Pt(10)
            
            # 设置背景色（灰色）
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F5F5F5')
            run._element.get_or_add_rPr().append(shading_elm)
    
    def process_element(self, element):
        """处理HTML元素"""
        if element.name == 'h1':
            self.add_heading(element.get_text(), 1)
        elif element.name == 'h2':
            self.add_heading(element.get_text(), 2)
        elif element.name == 'h3':
            self.add_heading(element.get_text(), 3)
        elif element.name == 'h4':
            self.add_heading(element.get_text(), 4)
        elif element.name == 'p':
            text = element.get_text()
            if text.strip():
                self.add_paragraph(text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                self.add_list_item(li.get_text())
        elif element.name == 'ol':
            for li in element.find_all('li'):
                self.add_list_item(li.get_text(), is_numbered=True)
        elif element.name == 'div' and 'code-block' in element.get('class', []):
            self.add_code_block(element.get_text())
        elif element.name in ['div', 'section']:
            # 递归处理子元素
            for child in element.children:
                if hasattr(child, 'name'):
                    self.process_element(child)
    
    def convert(self):
        """执行转换"""
        try:
            print(f"正在读取HTML文件: {self.html_file_path}")
            
            # 读取HTML文件
            with open(self.html_file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # 解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 添加文档标题
            title_element = soup.find('h1')
            if title_element:
                title_para = self.doc.add_heading(title_element.get_text(), 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 处理文档主体
            body = soup.find('body')
            if body:
                for element in body.children:
                    if hasattr(element, 'name') and element.name:
                        self.process_element(element)
            
            # 保存文档
            print(f"正在保存Word文档: {self.output_file_path}")
            self.doc.save(self.output_file_path)
            
            print(f"转换完成！文件已保存到: {self.output_file_path}")
            return True
            
        except Exception as e:
            print(f"转换过程中发生错误: {e}")
            return False


def main():
    """主函数"""
    # 默认文件路径
    html_file = r"g:\WeChat\WeChat Files\wxid_0izpcqrhi3m812\FileStorage\File\2025-07\video_capture\完整实习报告.html"
    
    # 检查文件是否存在
    if not os.path.exists(html_file):
        print(f"错误: HTML文件不存在: {html_file}")
        return False
    
    # 创建转换器实例
    converter = HTMLToWordConverter(html_file)
    
    # 执行转换
    success = converter.convert()
    
    if success:
        print("\n✅ HTML转Word转换成功完成！")
        print(f"📄 输出文件: {converter.output_file_path}")
        
        # 尝试打开输出文件
        try:
            import subprocess
            if sys.platform.startswith('win'):
                os.startfile(converter.output_file_path)
            elif sys.platform.startswith('darwin'):
                subprocess.call(['open', converter.output_file_path])
            else:
                subprocess.call(['xdg-open', converter.output_file_path])
            print("📂 已尝试打开输出文件")
        except Exception as e:
            print(f"⚠️  无法自动打开文件: {e}")
    else:
        print("\n❌ 转换失败")
        return False
    
    return True


if __name__ == "__main__":
    # 检查依赖包
    try:
        import docx
        import bs4
    except ImportError as e:
        print(f"❌ 缺少必要的依赖包: {e}")
        print("请运行以下命令安装依赖:")
        print("pip install python-docx beautifulsoup4 lxml")
        sys.exit(1)
    
    # 运行主程序
    main()
